from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
import os
import qrcode
import time
import subprocess
import zipfile

# Linux-compatible libraries
from pdf2docx import Converter
from PIL import Image, ImageDraw
from pdf2image import convert_from_path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document

app = Flask(__name__)

# 1. SETUP PATHS
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")
LOGO_PATH = os.path.join(BASE_DIR, "logo.png")

for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
    os.makedirs(folder, exist_ok=True)

POPPLER_PATH = None


def convert_via_libreoffice(input_path, output_folder):
    """Uses LibreOffice to convert Word/Excel/PPT to PDF on Linux."""
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf',
         '--outdir', output_folder, input_path],
        check=True,
        capture_output=True,
        text=True,
        timeout=60
    )
    return result


def get_libreoffice_output_path(input_path, output_folder):
    """
    LibreOffice names the output file based on the input filename stem.
    This helper finds the actual output path reliably.
    """
    base = os.path.splitext(os.path.basename(input_path))[0]
    expected = os.path.join(output_folder, base + ".pdf")
    return expected


def make_stylish_qr(data, logo_path, output_path):
    """Generates a QR code for the provided link data with an optional central logo."""
    qr = qrcode.QRCode(
        version=4,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white").convert("RGBA")

    try:
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            qr_width, qr_height = qr_img.size
            logo_max_size = qr_width // 3
            logo.thumbnail((logo_max_size, logo_max_size), Image.LANCZOS)

            bg_size = (logo.width + 20, logo.height + 20)
            logo_bg = Image.new("RGBA", bg_size, (255, 255, 255, 0))
            draw = ImageDraw.Draw(logo_bg)
            draw.rounded_rectangle([(0, 0), bg_size], radius=15, fill=(255, 255, 255, 255))

            pos_logo = ((bg_size[0] - logo.width) // 2, (bg_size[1] - logo.height) // 2)
            logo_bg.paste(logo, pos_logo, mask=logo)
            pos = ((qr_width - bg_size[0]) // 2, (qr_height - bg_size[1]) // 2)
            qr_img.paste(logo_bg, pos, mask=logo_bg)
    except Exception as e:
        print(f"QR Logo processing error (continuing without logo): {e}")

    qr_img.save(output_path)


@app.route('/')
def index():
    return render_template("index1.html")


@app.route('/upload')
def upload_page():
    return render_template("index2.html")


@app.route('/qr-converter')
def qr_upload_page():
    return render_template("index3.html")


@app.route('/api/convert', methods=['POST'])
def convert_api():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in request"}), 400

        files = request.files.getlist('file')
        conversion = request.args.get('type', '').upper().strip()

        if not files or files[0].filename == "":
            return jsonify({"error": "No file selected"}), 400

        unique_prefix = str(int(time.time()))
        output_path = ""

        # ────────────────────────────────────────────
        # ZIP  (multi-file support)
        # ────────────────────────────────────────────
        if "ZIP" in conversion:
            zip_filename = f"converted_{unique_prefix}.zip"
            output_path = os.path.join(OUTPUT_FOLDER, zip_filename)

            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file in files:
                    sf_name = secure_filename(file.filename)
                    temp_input = os.path.join(UPLOAD_FOLDER, f"zip_{unique_prefix}_{sf_name}")
                    file.save(temp_input)
                    zipf.write(temp_input, sf_name)
                    os.remove(temp_input)

        # ────────────────────────────────────────────
        # All single-file conversions
        # ────────────────────────────────────────────
        else:
            file = files[0]
            safe_name = f"{unique_prefix}_{secure_filename(file.filename)}"
            filename_no_ext = os.path.splitext(safe_name)[0]
            input_path = os.path.abspath(os.path.join(UPLOAD_FOLDER, safe_name))
            file.save(input_path)

            # ── PDF TO WORD ──────────────────────────
            if "PDF TO WORD" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + ".docx")
                cv = Converter(input_path)
                cv.convert(output_path)
                cv.close()

            # ── WORD / PPT / EXCEL → PDF (LibreOffice) ──
            elif any(x in conversion for x in ["WORD TO PDF", "PPT TO PDF", "EXCEL TO PDF"]):
                convert_via_libreoffice(input_path, OUTPUT_FOLDER)
                # LibreOffice uses the stem of the INPUT filename as the output name
                output_path = get_libreoffice_output_path(input_path, OUTPUT_FOLDER)
                if not os.path.exists(output_path):
                    raise Exception(
                        f"LibreOffice did not produce expected output: {output_path}"
                    )

            # ── PDF TO IMAGE ─────────────────────────
            elif "PDF TO IMAGE" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + ".jpg")
                images = convert_from_path(input_path, dpi=150, first_page=1, last_page=1)
                if images:
                    rgb_image = images[0].convert("RGB")
                    rgb_image.save(output_path, "JPEG", quality=90)
                else:
                    raise Exception("pdf2image returned no pages.")

            # ── IMAGE TO PDF ─────────────────────────
            elif "IMAGE TO PDF" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + ".pdf")
                img = Image.open(input_path)
                # Convert any mode to RGB (handles RGBA, P, L, etc.)
                if img.mode in ("RGBA", "P", "LA", "L"):
                    img = img.convert("RGB")
                # Fit image onto a letter-size page with margins
                page_w, page_h = letter  # points (72 dpi)
                margin = 40
                max_w = page_w - 2 * margin
                max_h = page_h - 2 * margin
                img_w, img_h = img.size
                scale = min(max_w / img_w, max_h / img_h)
                new_w = int(img_w * scale)
                new_h = int(img_h * scale)
                img = img.resize((new_w, new_h), Image.LANCZOS)

                c = canvas.Canvas(output_path, pagesize=letter)
                # Save resized image to a temp file for drawImage
                temp_img_path = os.path.join(UPLOAD_FOLDER, filename_no_ext + "_tmp.jpg")
                img.save(temp_img_path, "JPEG")
                x = (page_w - new_w) / 2
                y = (page_h - new_h) / 2
                c.drawImage(temp_img_path, x, y, width=new_w, height=new_h)
                c.save()
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)

            # ── TEXT TO PDF ──────────────────────────
            elif "TEXT TO PDF" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + ".pdf")

                # Read the text file with multiple encoding fallbacks
                content = ""
                for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                    try:
                        with open(input_path, "r", encoding=enc) as f:
                            content = f.read()
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue

                # Use ReportLab Platypus for proper word-wrap and multi-page support
                doc = SimpleDocTemplate(
                    output_path,
                    pagesize=letter,
                    leftMargin=inch,
                    rightMargin=inch,
                    topMargin=inch,
                    bottomMargin=inch,
                )
                styles = getSampleStyleSheet()
                normal_style = styles["Normal"]
                normal_style.fontName = "Helvetica"
                normal_style.fontSize = 11
                normal_style.leading = 16

                story = []
                for line in content.splitlines():
                    # Preserve blank lines as spacers
                    if line.strip() == "":
                        story.append(Spacer(1, 8))
                    else:
                        # Escape XML special characters for Paragraph
                        safe_line = (
                            line.replace("&", "&amp;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;")
                        )
                        story.append(Paragraph(safe_line, normal_style))

                if not story:
                    story.append(Paragraph("(empty file)", normal_style))

                doc.build(story)

            # ── WORD TO TEXT ─────────────────────────
            elif "WORD TO TEXT" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + ".txt")
                doc = Document(input_path)
                lines = []
                for para in doc.paragraphs:
                    lines.append(para.text)
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            lines.append(row_text)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write("\n".join(lines))

            # ── LINK TO QR CODE ──────────────────────
            elif "LINK TO QR CODE" in conversion or "QR" in conversion:
                output_path = os.path.join(OUTPUT_FOLDER, filename_no_ext + "_qr.png")

                # The front-end sends the URL as a text/plain Blob named "link.txt"
                with open(input_path, "r", errors="ignore") as f:
                    link_content = f.read().strip()

                qr_link = link_content if link_content else "https://google.com"

                logo_file = request.files.get('logo')
                current_logo_path = LOGO_PATH
                temp_logo_created = False

                if logo_file and logo_file.filename != "":
                    temp_logo_name = f"tmp_{unique_prefix}_{secure_filename(logo_file.filename)}"
                    temp_logo_path = os.path.join(UPLOAD_FOLDER, temp_logo_name)
                    logo_file.save(temp_logo_path)
                    current_logo_path = temp_logo_path
                    temp_logo_created = True

                make_stylish_qr(qr_link, current_logo_path, output_path)

                if temp_logo_created and os.path.exists(current_logo_path):
                    os.remove(current_logo_path)

            else:
                return jsonify({"error": f"Unknown conversion type: '{conversion}'"}), 400

        if not output_path or not os.path.exists(output_path):
            return jsonify({"error": f"Output file was not generated for: {conversion}"}), 500

        return send_file(output_path, as_attachment=True)

    except subprocess.CalledProcessError as e:
        err_msg = e.stderr if hasattr(e, 'stderr') and e.stderr else str(e)
        return jsonify({"error": f"LibreOffice conversion failed: {err_msg}"}), 500
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
